import os
import shutil
from docx import Document
from openpyxl import Workbook

def create_demo_data():
    base_dir = os.path.join(os.getcwd(), "demo_data")
    if os.path.exists(base_dir):
        shutil.rmtree(base_dir)
    os.makedirs(base_dir)
    
    print(f"Creating demo files in: {base_dir}")
    
    # 1. Procedimiento (Word)
    doc_proc = Document()
    doc_proc.add_heading('Manual de Procedimientos Operativos', 0)
    doc_proc.add_paragraph('Este documento describe el procedimiento para...')
    doc_proc.save(os.path.join(base_dir, "borrador_procedimiento.docx"))
    
    # 2. Formato (Excel)
    wb_form = Workbook()
    ws = wb_form.active
    ws['A1'] = "FORMATO DE INSPECCION"
    ws['A2'] = "Fecha"
    ws['B2'] = "Inspector"
    wb_form.save(os.path.join(base_dir, "plantilla_inspeccion.xlsx"))
    
    # 3. Acta (Word)
    doc_acta = Document()
    doc_acta.add_heading('Acta de Reunión Semanal', 0)
    doc_acta.add_paragraph('Minuta de la reunión celebrada el día...')
    doc_acta.save(os.path.join(base_dir, "notas_reunion_ayer.docx"))
    
    # 4. Create Duplicate (Word)
    # Copy the procedure file
    shutil.copy(os.path.join(base_dir, "borrador_procedimiento.docx"), 
                os.path.join(base_dir, "copia_procedimiento.docx"))
                
    print("Files created:")
    print("- borrador_procedimiento.docx (Topic: PROCEDIMIENTO)")
    print("- plantilla_inspeccion.xlsx (Topic: FORMATO)")
    print("- notas_reunion_ayer.docx (Topic: ACTA)")
    print("- copia_procedimiento.docx (DUPLICATE of borrador_procedimiento)")

if __name__ == "__main__":
    create_demo_data()
